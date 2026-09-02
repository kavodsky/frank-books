"""Docx template styles and pandoc smoke (roadmap 6.1, 6.3)."""

from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

import pytest
from docx import Document
from docx.shared import RGBColor

from frank.application.render_book import RenderPorts, render_book
from frank.domain.model.annotation import Annotation, Morphology, SenseUnit, Token
from frank.domain.model.book import (
    Book,
    BookStatus,
    BookStructure,
    Chapter,
    Paragraph,
    ParagraphStatus,
    Passage,
    Sentence,
)
from frank.domain.model.frank import (
    FrankRecord,
    ModelTier,
    ParagraphOutput,
    SenseUnitTranslation,
    WordNote,
)
from frank.domain.model.layout import (
    LayoutBook,
    LayoutChapter,
    LayoutParagraph,
    LayoutPassage,
    LayoutRun,
    RunStyle,
)
from frank.infrastructure.persistence.repositories import SqliteBookRepository
from frank.infrastructure.persistence.tables import create_book_db
from frank.infrastructure.rendering.docx_renderer import write_docx, write_template

_GREEN = RGBColor(0x2E, 0x7D, 0x32)
_REPO = Path(__file__).resolve().parents[2]
_TEMPLATE = _REPO / "templates" / "frank.docx"


def _layout() -> LayoutBook:
    adapted = LayoutParagraph(
        runs=(
            LayoutRun(text="Oliver kommt", style=RunStyle.ORIGINAL),
            LayoutRun(text=" (", style=RunStyle.TRANSLATION),
            LayoutRun(text="Олівер іде", style=RunStyle.TRANSLATION),
            LayoutRun(text="; Oliver – Олівер", style=RunStyle.GLOSS),
            LayoutRun(text=")", style=RunStyle.TRANSLATION),
            LayoutRun(text=".", style=RunStyle.ORIGINAL),
        )
    )
    unadapted = LayoutParagraph(
        runs=(LayoutRun(text="Oliver kommt.", style=RunStyle.UNADAPTED),)
    )
    return LayoutBook(
        title="Oliver Twist",
        author="Charles Dickens",
        source_url="https://example.test/oliver",
        license_note="public domain",
        chapters=(
            LayoutChapter(
                title="I",
                passages=(LayoutPassage(adapted=(adapted,), unadapted=(unadapted,)),),
            ),
        ),
        marker="— згенеровано до пасажу 1 —",
    )


@pytest.mark.integration
def test_template_defines_frank_character_styles() -> None:
    assert _TEMPLATE.is_file()
    styles = {item.name for item in Document(str(_TEMPLATE)).styles}
    assert {
        "FrankOriginal",
        "FrankTranslation",
        "FrankGloss",
        "FrankNote",
        "FrankUnadapted",
    } <= styles


@pytest.mark.integration
def test_docx_uses_green_translation_and_marker(tmp_path) -> None:
    path = tmp_path / "oliver.docx"
    write_docx(_layout(), path)
    document = Document(str(path))
    texts = [para.text for para in document.paragraphs]
    assert "Oliver Twist" in texts[0]
    assert "Charles Dickens" in texts[1]
    assert any("Олівер іде" in item for item in texts)
    assert texts[-1] == "— згенеровано до пасажу 1 —"
    assert document.styles["FrankTranslation"].font.color.rgb == _GREEN
    assert document.styles["FrankGloss"].font.italic is True
    names = {
        run.style.name for para in document.paragraphs for run in para.runs if run.style
    }
    assert "FrankTranslation" in names
    assert "FrankOriginal" in names


@pytest.mark.integration
def test_docx_converts_with_pandoc(tmp_path) -> None:
    pandoc = shutil.which("pandoc")
    if pandoc is None:
        pytest.skip("pandoc not installed")
    path = tmp_path / "oliver.docx"
    write_docx(_layout(), path)
    html = tmp_path / "oliver.html"
    epub = tmp_path / "oliver.epub"
    subprocess.run([pandoc, str(path), "-o", str(html)], check=True)
    subprocess.run([pandoc, str(path), "-o", str(epub)], check=True)
    assert html.is_file() and html.stat().st_size > 0
    assert epub.is_file() and epub.stat().st_size > 0


@pytest.mark.integration
def test_write_template_is_idempotent(tmp_path) -> None:
    path = tmp_path / "frank.docx"
    write_template(path)
    write_template(path)
    names = {item.name for item in Document(str(path)).styles}
    assert "FrankTranslation" in names


@pytest.mark.integration
def test_render_book_writes_completed_passages(tmp_path) -> None:
    engine = create_book_db(tmp_path / "book.db")
    books = SqliteBookRepository(engine)
    _seed_complete(books)
    path = tmp_path / "oliver.docx"
    report = render_book(
        RenderPorts(
            open_books=lambda _slug: books,
            open_records=lambda _slug: books,
            write_docx=write_docx,
        ),
        "oliver-de",
        path,
    )
    assert report.passages == 1
    document = Document(str(path))
    joined = "\n".join(para.text for para in document.paragraphs)
    assert "Oliver kommt (Олівер іде; Oliver – Олівер)." in joined
    assert "— згенеровано до пасажу 1 —" in joined


def _seed_complete(books: SqliteBookRepository) -> None:
    structure = BookStructure(
        book=Book(
            id="b",
            slug="oliver-de",
            lang="de",
            title="Oliver Twist",
            author="Charles Dickens",
            source_url="https://example.test/oliver",
            license_note="public domain",
            status=BookStatus.INGESTED,
        ),
        chapters=(Chapter(id="c1", book_id="b", index=1, title="I"),),
        passages=(Passage(id="pass-1", chapter_id="c1", index=1),),
        paragraphs=(
            Paragraph(
                id="p1",
                chapter_id="c1",
                passage_id="pass-1",
                index=1,
                raw_text="Oliver kommt.",
                hash="h1",
                status=ParagraphStatus.COMPLETE,
            ),
        ),
    )
    books.save_structure(structure)
    books.replace_passages("oliver-de", structure)
    books.replace_annotation(
        "oliver-de",
        Annotation(
            sentences=(
                Sentence(id="s1", paragraph_id="p1", index=1, text="Oliver kommt."),
            ),
            tokens=(
                Token(
                    id="t1",
                    sentence_id="s1",
                    index=1,
                    surface="Oliver",
                    lemma="Oliver",
                    upos="PROPN",
                    morph=Morphology(),
                ),
                Token(
                    id="t2",
                    sentence_id="s1",
                    index=2,
                    surface="kommt",
                    lemma="kommen",
                    upos="VERB",
                    morph=Morphology(),
                ),
                Token(
                    id="t3",
                    sentence_id="s1",
                    index=3,
                    surface=".",
                    lemma=".",
                    upos="PUNCT",
                    morph=Morphology(),
                ),
            ),
            sense_units=(
                SenseUnit(
                    id="u1",
                    sentence_id="s1",
                    index=1,
                    start_index=1,
                    end_index=3,
                ),
            ),
        ),
    )
    books.save_paragraph_output(
        "oliver-de",
        ParagraphOutput(
            paragraph_id="p1",
            records=(
                FrankRecord(
                    sentence_id="s1",
                    units=(
                        SenseUnitTranslation(
                            source_span=(1, 3),
                            natural_uk="Олівер іде",
                        ),
                    ),
                    idiomatic_uk="Олівер іде.",
                    word_notes=(
                        WordNote(
                            surface="Oliver",
                            lemma="Oliver",
                            morph_note_uk="",
                            gloss_uk="Олівер",
                        ),
                    ),
                    tier=ModelTier.FAST,
                ),
            ),
            status=ParagraphStatus.COMPLETE,
        ),
    )
