"""python-docx adapter: LayoutBook to templates/frank.docx styles (roadmap 6.1–6.2)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.text.run import Run
from lxml.etree import _Element

from frank.domain.model.layout import (
    LayoutBook,
    LayoutChapter,
    LayoutParagraph,
    LayoutPassage,
    LayoutRun,
    RunStyle,
)

_GREEN = RGBColor(0x2E, 0x7D, 0x32)
_BLACK = RGBColor(0x00, 0x00, 0x00)
_FONT = "Georgia"
_STYLES = {
    RunStyle.ORIGINAL: "FrankOriginal",
    RunStyle.TRANSLATION: "FrankTranslation",
    RunStyle.GLOSS: "FrankGloss",
    RunStyle.NOTE: "FrankNote",
    RunStyle.UNADAPTED: "FrankUnadapted",
}


def write_docx(layout: LayoutBook, path: Path) -> None:
    document = Document(str(_template_path()))
    _clear_body(document)
    _title_page(document, layout)
    for chapter in layout.chapters:
        _chapter(document, chapter)
    _marker(document, layout.marker)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


def write_template(path: Path) -> None:
    """Create the named-style sheet (roadmap 6.1). Call when the template is missing."""
    document = Document()
    _set_normal(document)
    _set_heading(document)
    _add_char(document, "FrankOriginal", _BLACK)
    _add_char(document, "FrankTranslation", _GREEN)
    _add_italic_char(document, "FrankGloss", _GREEN)
    _add_italic_char(document, "FrankNote", _GREEN)
    _add_char(document, "FrankUnadapted", _BLACK)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


def _clear_body(document: Document) -> None:
    body = document.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _title_page(document: Document, layout: LayoutBook) -> None:
    _block(document, layout.title, WD_ALIGN_PARAGRAPH.CENTER, Pt(22))
    if layout.author:
        _block(document, layout.author, WD_ALIGN_PARAGRAPH.CENTER, Pt(14))
    note = layout.license_note.strip()
    if layout.source_url.strip():
        note = f"{note}\n{layout.source_url}".strip()
    if note:
        _block(document, note, WD_ALIGN_PARAGRAPH.LEFT, Pt(11))


def _chapter(document: Document, chapter: LayoutChapter) -> None:
    heading = document.add_paragraph(chapter.title, style="Heading 1")
    _widow(heading)
    for passage in chapter.passages:
        _passage(document, passage)


def _passage(document: Document, passage: LayoutPassage) -> None:
    for paragraph in passage.adapted:
        _body(document, paragraph)
    document.add_paragraph("")
    for paragraph in passage.unadapted:
        _body(document, paragraph)
    document.add_paragraph("")


def _body(document: Document, paragraph: LayoutParagraph) -> None:
    docx_para = document.add_paragraph()
    docx_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _widow(docx_para)
    for item in paragraph.runs:
        _run(docx_para, item)


def _run(paragraph: DocxParagraph, item: LayoutRun) -> None:
    run = paragraph.add_run(item.text)
    run.style = _STYLES[item.style]


def _marker(document: Document, text: str) -> None:
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _widow(para)
    run = para.add_run(text)
    run.italic = True


def _block(
    document: Document,
    text: str,
    align: WD_ALIGN_PARAGRAPH,
    size: Pt,
) -> None:
    para = document.add_paragraph()
    para.alignment = align
    _widow(para)
    run = para.add_run(text)
    run.font.size = size
    run.font.name = _FONT
    _run_fonts(run)


def _set_normal(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(12)
    _rfonts(style.font._element, _FONT)
    fmt = style.paragraph_format
    fmt.line_spacing = 1.35
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.widow_control = True


def _set_heading(document: Document) -> None:
    style = document.styles["Heading 1"]
    style.font.name = _FONT
    style.font.size = Pt(16)
    style.font.color.rgb = _BLACK
    _rfonts(style.font._element, _FONT)
    style.paragraph_format.page_break_before = True
    style.paragraph_format.widow_control = True


def _add_char(document: Document, name: str, color: RGBColor) -> None:
    style = document.styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
    style.font.name = _FONT
    style.font.size = Pt(12)
    style.font.color.rgb = color
    _rfonts(style.font._element, _FONT)


def _add_italic_char(document: Document, name: str, color: RGBColor) -> None:
    _add_char(document, name, color)
    document.styles[name].font.italic = True


def _widow(paragraph: DocxParagraph) -> None:
    paragraph.paragraph_format.widow_control = True


def _run_fonts(run: Run) -> None:
    _rfonts(run._element, _FONT)


def _rfonts(element: _Element, name: str) -> None:
    props = element.find(qn("w:rPr"))
    if props is None:
        props = element
    fonts = props.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        props.insert(0, fonts)
    for key in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(key), name)


def _template_path() -> Path:
    bundled = Path(__file__).resolve().parent / "frank.docx"
    if bundled.is_file():
        return bundled
    return Path(__file__).resolve().parents[3] / "templates" / "frank.docx"
